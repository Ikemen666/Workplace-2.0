
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Style & Formatting Constants (per Monet SOUL) ---
# NOTE: This template uses the Monet SOUL standard: Aptos Display 16/12pt.
# The SKILL.md for minute-creator provides the context for these values.
FONT_NAME = 'Aptos Display'
TITLE_PT = 16
CONTEXT_PT = 12
BODY_PT = 12
LINE_SPACING = 1.0
COLOR_BLACK = RGBColor(0, 0, 0)

# --- Document Setup ---
doc = Document()
# Set default styles for the document
style = doc.styles['Normal']
style.font.name = FONT_NAME
style.font.color.rgb = COLOR_BLACK
# Set East Asian font compatibility
r = style.element.rPr
r.rFonts.set(qn('w:eastAsia'), FONT_NAME)
# Set default paragraph line spacing
style.paragraph_format.line_spacing = LINE_SPACING


# --- Helper Functions (Corrected & Robust Version) ---

def apply_run_style(run, font_name=FONT_NAME, size_pt=BODY_PT, bold=False,
                    italic=False, color=COLOR_BLACK):
    """Applies font-related styles (font, size, bold, etc.) to a Run object."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    # Set font for East Asian characters to ensure compatibility
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts_list = rPr.xpath('./w:rFonts')
    if rFonts_list:
        rFonts_list[0].set(qn('w:eastAsia'), font_name)
    else:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.insert(0, rFonts)

def apply_paragraph_style(paragraph):
    """Applies paragraph-level styles (e.g., line spacing)."""
    paragraph.paragraph_format.line_spacing = LINE_SPACING

def add_title(doc, text):
    """Adds a centered, styled title using a standard paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_run_style(run, size_pt=TITLE_PT, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_paragraph_style(p)
    
def add_l1_heading(doc, text):
    """Adds a styled Level 1 heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_run_style(run, bold=True)
    apply_paragraph_style(p)

def add_context_intro(doc, text):
    """Adds an italicized context paragraph below a heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_run_style(run, size_pt=CONTEXT_PT, italic=True)
    apply_paragraph_style(p)

def add_bullet(doc, lead_phrase, detail_text, level=1):
    """Adds a styled bullet point with a bolded lead phrase."""
    # Mapping of levels to markers and indents
    markers = {1: '●', 2: '1.', 3: 'a.', 4: '○'}
    indent_pt = 20 * level
    
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(indent_pt)
    
    # Add marker, bold lead phrase, and detail text as separate runs
    marker_run = p.add_run(f"{markers.get(level, '●')} ")
    apply_run_style(marker_run)

    lead_run = p.add_run(f"{lead_phrase}：")
    apply_run_style(lead_run, bold=True)

    detail_run = p.add_run(detail_text)
    apply_run_style(detail_run)
    
    # Apply paragraph-level styles
    apply_paragraph_style(p)

def add_action_item(doc, timeframe, owner_summary, task_desc):
    """Adds a formatted 'Next Step' action item."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)

    # Marker
    marker_run = p.add_run("● ")
    apply_run_style(marker_run)
    
    # Timeframe (Bold)
    tf_run = p.add_run(timeframe)
    apply_run_style(tf_run, bold=True)

    # Owner/Summary (Underlined)
    owner_run = p.add_run(f" {owner_summary}")
    owner_run.underline = True
    apply_run_style(owner_run)
    
    # Task Description
    detail_run = p.add_run(f"：{task_desc}")
    apply_run_style(detail_run)

    apply_paragraph_style(p)

# --- Main Execution ---
if __name__ == "__main__":
    # This is a template. The agent should dynamically fill the content.
    # The following is an example structure.

    # 1. Title
    add_title(doc, "[YYMMDD] Project Meeting Memo – [会议名称]")

    # 2. Section I
    add_l1_heading(doc, "I. [主节标题]")
    add_context_intro(doc, "[一句话点出本节核心背景或结论。]")
    add_bullet(doc, "要点一", "详细叙述...")
    add_bullet(doc, "要点二", "详细叙述...")

    # 3. Section II
    add_l1_heading(doc, "II. [主节标题]")
    add_bullet(doc, "要点一", "详细叙述...")

    # 4. Open Items (if any)
    add_l1_heading(doc, "III. 争议与未决事项")
    add_bullet(doc, "争议点", "[A 主张 X] vs [B 主张 Y] | 当前状态：[未决]")
    
    # 5. Next Steps (always last)
    add_l1_heading(doc, "IV. 下一步计划")
    add_action_item("**[Timeframe]**", "[Owner/Team_Summary]", "[动词开头的任务描述]")
    add_action_item("**TBD**", "[Owner/Team_Summary]", "[动词开头的任务描述]")

    # --- Save Document ---
    save_dir = os.path.expanduser("~/Desktop/Meeting Minutes")
    os.makedirs(save_dir, exist_ok=True)
    file_name = "YYYYMMDD_Project_Memo_Template.docx"
    save_path = os.path.join(save_dir, file_name)
    doc.save(save_path)

    print(f"Template memo generated and saved to: {save_path}")
